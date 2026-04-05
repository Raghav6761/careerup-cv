import io
import re
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER
from bidi.algorithm import get_display
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


FONT_DIR = os.path.join(os.path.dirname(__file__), "fonts")

_EMPTY_PLACEHOLDERS = [
    "לא צוין", "לא צויין", "לא צויינו", "לא צוינו", "לא סופק", "לא מולא", "לא קיים",
    "לא רלוונטי", "לא קיים במקור", "אין", "ללא", "—", "-",
    "not specified", "not provided", "n/a", "none", "not available",
    "no data", "not applicable", "na", "tbd", "n.a.", "n.a",
]

# Matches any Hebrew "לא + word" phrase (covers all verb-form variants the AI may produce).
# Intentionally broad: any sentence starting with "לא <non-space>" is treated as a
# "not specified" placeholder. This is safe in CV context — no real section content
# begins with "לא" in the improve/build flows.
_HEBREW_NEGATION_RE = re.compile(r"^לא\s+\S", re.UNICODE)

def _is_empty_content(text: str) -> bool:
    if not text or not text.strip():
        return True
    cleaned = text.strip().rstrip(".").strip()
    cleaned_lower = cleaned.lower()
    if cleaned_lower in _EMPTY_PLACEHOLDERS:
        return True
    # Fast regex catch: any line that starts with "לא <word>" in Hebrew is a placeholder
    if _HEBREW_NEGATION_RE.match(cleaned):
        return True
    for placeholder in _EMPTY_PLACEHOLDERS:
        if (cleaned_lower.startswith(placeholder.lower() + " ")
                or cleaned_lower.startswith(placeholder.lower() + ":")
                or cleaned_lower.startswith(placeholder.lower() + "-")):
            return True
    return False

def _filter_list(items: list) -> list:
    if not items:
        return []
    return [item for item in items if item and not _is_empty_content(str(item))]


_PDF_COMPRESSION_LEVELS = [
    {"margin_mm": 15, "top_mm": 10, "bot_mm": 10, "sec_before": 8,   "body_lead": 14,  "bul_lead": 13,  "body_size": 10,  "sec_size": 13, "contact_after": 6, "name_lead": 26},
    {"margin_mm": 13, "top_mm": 9,  "bot_mm": 9,  "sec_before": 6,   "body_lead": 13,  "bul_lead": 12,  "body_size": 9,   "sec_size": 12, "contact_after": 5, "name_lead": 24},
    {"margin_mm": 11, "top_mm": 7,  "bot_mm": 7,  "sec_before": 4,   "body_lead": 11,  "bul_lead": 10,  "body_size": 8,   "sec_size": 10, "contact_after": 3, "name_lead": 20},
]


def _count_pdf_pages(pdf_bytes: bytes) -> int:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            return len(pdf.pages)
    except Exception:
        return 1

def _has_real_exp(exp: dict) -> bool:
    title = exp.get("title", "").strip()
    company = exp.get("company", "").strip()
    period = exp.get("period", "").strip()
    if (title and not _is_empty_content(title)) or (company and not _is_empty_content(company)) or (period and not _is_empty_content(period)):
        return True
    achs = [a for a in exp.get("achievements", []) if a.strip() and not _is_empty_content(a)]
    if achs:
        return True
    return False

def _has_real_edu(edu: dict) -> bool:
    degree = edu.get("degree", "").strip()
    institution = edu.get("institution", "").strip()
    year = edu.get("year", "").strip()
    return bool(
        (degree and not _is_empty_content(degree)) or
        (institution and not _is_empty_content(institution)) or
        (year and not _is_empty_content(year))
    )


def register_hebrew_font():
    font_path = os.path.join(FONT_DIR, "Assistant-Regular.ttf")
    font_bold_path = os.path.join(FONT_DIR, "Assistant-Bold.ttf")

    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont("Assistant", font_path))
        if os.path.exists(font_bold_path):
            pdfmetrics.registerFont(TTFont("Assistant-Bold", font_bold_path))
        else:
            pdfmetrics.registerFont(TTFont("Assistant-Bold", font_path))
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        registerFontFamily("Assistant", normal="Assistant", bold="Assistant-Bold")
        return "Assistant"
    else:
        return "Helvetica"


def _clean_hebrew_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\u200f', '').replace('\u200e', '')
    text = text.replace('(', ' - ').replace(')', '').replace('（', ' - ').replace('）', '')
    text = text.replace('"', "'").replace('"', "'").replace('"', "'")
    text = text.replace('״', "'").replace('׳', "'")
    return text


def reshape_hebrew(text: str) -> str:
    if not text:
        return ""
    text = _clean_hebrew_text(text)
    lines = text.split('\n')
    result_lines = []
    for line in lines:
        if not line.strip():
            result_lines.append(line)
            continue
        displayed = str(get_display(line, base_dir='R'))
        result_lines.append(displayed)
    return '\n'.join(result_lines)


def reshape_hebrew_paragraph(text: str, font_name="Assistant", font_size=9, max_width=None) -> str:
    if not text:
        return ""
    text = _clean_hebrew_text(text)
    text = text.replace('\n', ' ')
    if max_width is None:
        max_width = (210 - 2 * 20) * mm
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = text.split()
    lines = []
    current_words = []
    current_width = 0
    space_w = stringWidth(' ', font_name, font_size)
    for word in words:
        w = stringWidth(word, font_name, font_size)
        test_width = current_width + (space_w if current_words else 0) + w
        if current_words and test_width > max_width:
            line_text = ' '.join(current_words)
            lines.append(str(get_display(line_text, base_dir='R')))
            current_words = [word]
            current_width = w
        else:
            current_words.append(word)
            current_width = test_width
    if current_words:
        line_text = ' '.join(current_words)
        lines.append(str(get_display(line_text, base_dir='R')))
    return '<br/>'.join(lines)


def _make_section_separator(width_mm=170):
    return HRFlowable(
        width="100%",
        thickness=1,
        color=HexColor("#7fb3d8"),
        spaceAfter=2,
        spaceBefore=6
    )


def _make_job_separator(width_mm=170):
    return HRFlowable(
        width="100%",
        thickness=1,
        color=HexColor("#e2e8f0"),
        spaceAfter=2,
        spaceBefore=3
    )


def _get_pdf_styles(font_name, bold_font, cparams=None):
    if cparams is None:
        cparams = _PDF_COMPRESSION_LEVELS[0]
    bs = cparams["body_size"]
    ss = cparams["sec_size"]
    bl = cparams["body_lead"]
    ul = cparams["bul_lead"]
    sb = cparams["sec_before"]
    ca = cparams["contact_after"]
    nl = cparams["name_lead"]
    return {
        "name": ParagraphStyle(
            "Name",
            fontName=bold_font,
            fontSize=20,
            leading=nl,
            alignment=TA_CENTER,
            textColor=HexColor("#2c3e50"),
            spaceAfter=2
        ),
        "contact": ParagraphStyle(
            "Contact",
            fontName=font_name,
            fontSize=bs,
            leading=ul,
            alignment=TA_CENTER,
            textColor=HexColor("#555555"),
            spaceAfter=ca
        ),
        "section_header": ParagraphStyle(
            "SectionHeader",
            fontName=bold_font,
            fontSize=ss,
            leading=ss + 4,
            alignment=TA_RIGHT,
            textColor=HexColor("#2c3e50"),
            spaceAfter=2,
            spaceBefore=sb
        ),
        "job_title": ParagraphStyle(
            "JobTitle",
            fontName=bold_font,
            fontSize=bs,
            leading=bl,
            alignment=TA_RIGHT,
            textColor=HexColor("#333333"),
            spaceAfter=1,
            spaceBefore=max(int(sb / 2), 2)
        ),
        "job_period": ParagraphStyle(
            "JobPeriod",
            fontName=font_name,
            fontSize=bs,
            leading=ul,
            alignment=TA_RIGHT,
            textColor=HexColor("#6b7c93"),
            spaceAfter=1
        ),
        "body": ParagraphStyle(
            "Body",
            fontName=font_name,
            fontSize=bs,
            leading=bl,
            alignment=TA_RIGHT,
            textColor=HexColor("#333333"),
            spaceAfter=2
        ),
        "body_bold": ParagraphStyle(
            "BodyBold",
            fontName=bold_font,
            fontSize=bs,
            leading=bl,
            alignment=TA_RIGHT,
            textColor=HexColor("#2c3e50"),
            spaceAfter=1
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            fontName=font_name,
            fontSize=bs,
            leading=ul,
            alignment=TA_RIGHT,
            textColor=HexColor("#333333"),
            spaceAfter=1,
            rightIndent=8
        ),
    }


def export_cv_to_pdf(cv_data: dict, max_pages: int = 1) -> bytes:
    font_name = register_hebrew_font()
    bold_font = f"{font_name}-Bold" if font_name == "Assistant" else "Helvetica-Bold"

    for level in range(len(_PDF_COMPRESSION_LEVELS)):
        cparams = _PDF_COMPRESSION_LEVELS[level]
        margin_mm = cparams["margin_mm"]

        buffer = io.BytesIO()
        top_mm = cparams.get("top_mm", 12)
        bot_mm = cparams.get("bot_mm", 12)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=margin_mm * mm,
            leftMargin=margin_mm * mm,
            topMargin=top_mm * mm,
            bottomMargin=bot_mm * mm
        )
        styles = _get_pdf_styles(font_name, bold_font, cparams)
        elements = []

        name = cv_data.get("full_name", "")
        if name:
            elements.append(Paragraph(reshape_hebrew(name), styles["name"]))

        contact = cv_data.get("contact", {})
        contact_parts = []
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("city"):
            contact_parts.append(reshape_hebrew(contact["city"]))
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact_parts:
            elements.append(Paragraph(" | ".join(contact_parts), styles["contact"]))

        elements.append(HRFlowable(width="100%", thickness=2, color=HexColor("#2c3e50"), spaceAfter=4, spaceBefore=2))

        summary = cv_data.get("professional_summary", "")
        if summary and not _is_empty_content(summary):
            elements.append(Paragraph(reshape_hebrew("תקציר מקצועי"), styles["section_header"]))
            elements.append(_make_section_separator())
            elements.append(Paragraph(reshape_hebrew_paragraph(summary), styles["body"]))

        experience = [e for e in cv_data.get("experience", []) if _has_real_exp(e)]
        if experience:
            elements.append(Paragraph(reshape_hebrew("ניסיון תעסוקתי"), styles["section_header"]))
            elements.append(_make_section_separator())
            for idx, exp in enumerate(experience):
                title = exp.get("title", "")
                company = exp.get("company", "")
                period = exp.get("period", "")
                if idx > 0:
                    elements.append(_make_job_separator())
                header_parts = []
                if period:
                    header_parts.append(period)
                if title:
                    header_parts.append(title)
                if company:
                    header_parts.append(company)
                if header_parts:
                    elements.append(Paragraph(reshape_hebrew(" | ".join(header_parts)), styles["job_title"]))
                for ach in exp.get("achievements", []):
                    if ach.strip() and not _is_empty_content(ach):
                        elements.append(Paragraph(reshape_hebrew(f"• {ach}"), styles["bullet"]))
                honors = exp.get("honors", "")
                if honors and honors.strip() and not _is_empty_content(honors):
                    elements.append(Paragraph(reshape_hebrew(f"★ {honors}"), styles["bullet"]))

        education = [e for e in cv_data.get("education", []) if _has_real_edu(e)]
        if education:
            elements.append(Paragraph(reshape_hebrew("השכלה"), styles["section_header"]))
            elements.append(_make_section_separator())
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")
                honors = edu.get("honors", "")
                parts = []
                if year:
                    parts.append(year)
                if degree:
                    parts.append(degree)
                if institution:
                    parts.append(institution)
                if parts:
                    elements.append(Paragraph(reshape_hebrew(" | ".join(parts)), styles["body"]))
                if honors and honors.strip() and not _is_empty_content(honors):
                    elements.append(Paragraph(reshape_hebrew(f"★ {honors}"), styles["bullet"]))

        skills = cv_data.get("skills", {})
        technical = _filter_list(skills.get("technical", []))
        soft = _filter_list(skills.get("soft", []))
        if technical or soft:
            elements.append(Paragraph(reshape_hebrew("מיומנויות"), styles["section_header"]))
            elements.append(_make_section_separator())
            if technical:
                elements.append(Paragraph(reshape_hebrew(", ".join(technical)), styles["body"]))
            if soft:
                elements.append(Paragraph(reshape_hebrew(", ".join(soft)), styles["body"]))

        lang_parts = []
        for lang in cv_data.get("languages", []):
            lang_name = lang.get("language", "").strip()
            lang_level = lang.get("level", "").strip()
            if lang_name and not _is_empty_content(lang_name):
                part = lang_name
                if lang_level and not _is_empty_content(lang_level):
                    part += f" – {lang_level}"
                lang_parts.append(part)
        if lang_parts:
            elements.append(Paragraph(reshape_hebrew("שפות"), styles["section_header"]))
            elements.append(_make_section_separator())
            elements.append(Paragraph(reshape_hebrew(" | ".join(lang_parts)), styles["body"]))

        military = _filter_list(cv_data.get("military", []))
        if military:
            elements.append(Paragraph(reshape_hebrew("שירות צבאי / לאומי"), styles["section_header"]))
            elements.append(_make_section_separator())
            for item in military:
                elements.append(Paragraph(reshape_hebrew(f"• {item}"), styles["bullet"]))

        volunteering = _filter_list(cv_data.get("volunteering", []))
        if volunteering:
            elements.append(Paragraph(reshape_hebrew("התנדבות"), styles["section_header"]))
            elements.append(_make_section_separator())
            for item in volunteering:
                elements.append(Paragraph(reshape_hebrew(f"• {item}"), styles["bullet"]))

        projects = _filter_list(cv_data.get("projects", []))
        if projects:
            elements.append(Paragraph(reshape_hebrew("פרויקטים עצמאיים"), styles["section_header"]))
            elements.append(_make_section_separator())
            for item in projects:
                elements.append(Paragraph(reshape_hebrew(f"• {item}"), styles["bullet"]))

        additional = _filter_list(cv_data.get("additional", []))
        if additional:
            elements.append(Paragraph(reshape_hebrew("מידע נוסף"), styles["section_header"]))
            elements.append(_make_section_separator())
            for item in additional:
                elements.append(Paragraph(reshape_hebrew(f"• {item}"), styles["bullet"]))

        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        if _count_pdf_pages(pdf_bytes) <= max_pages or level == len(_PDF_COMPRESSION_LEVELS) - 1:
            return pdf_bytes

    return buffer.getvalue()


def _add_docx_separator_line(doc, color='7fb3d8', size='8', space_before=3, space_after=1):
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(space_before)
    sep.paragraph_format.space_after = Pt(space_after)
    sep_pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    sep_pPr.append(pBdr)


def _add_docx_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)

    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    bCs = OxmlElement('w:bCs')
    rPr.append(bCs)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), 'Assistant')

    _add_docx_separator_line(doc, '7fb3d8', '8', 0, 1)


def _add_docx_body_paragraph(doc, text, is_rtl=True, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    if bold:
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    if indent:
        if is_rtl:
            p.paragraph_format.right_indent = Pt(indent)
        else:
            p.paragraph_format.left_indent = Pt(indent)
    if is_rtl:
        _set_docx_rtl(p)
    return p


def _add_docx_bullet_paragraph(doc, text, is_rtl=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(12)
    if is_rtl:
        p.paragraph_format.right_indent = Pt(8)
        _set_docx_rtl(p)
    else:
        p.paragraph_format.left_indent = Pt(8)
    return p


def _add_docx_job_header(doc, text, is_rtl=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = Pt(13)
    if is_rtl:
        _set_docx_rtl(p)
    return p


def _add_docx_hr(doc):
    _add_docx_separator_line(doc, '2c3e50', '16', 2, 4)


def _add_docx_job_separator(doc):
    _add_docx_separator_line(doc, 'e2e8f0', '8', 3, 2)


def export_cv_to_docx(cv_data: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Assistant"
    font.size = Pt(9)
    font.color.rgb = RGBColor(51, 51, 51)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(1)

    for section in doc.sections:
        section.right_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    name = cv_data.get("full_name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(24)
        _set_docx_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = cv_data.get("contact", {})
    contact_parts = []
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("city"):
        contact_parts.append(contact["city"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(12)
        _set_docx_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_docx_hr(doc)

    summary = cv_data.get("professional_summary", "")
    if summary and not _is_empty_content(summary):
        _add_docx_section_header(doc, "תקציר מקצועי")
        _add_docx_body_paragraph(doc, summary)

    experience = [e for e in cv_data.get("experience", []) if _has_real_exp(e)]
    if experience:
        _add_docx_section_header(doc, "ניסיון תעסוקתי")
        for idx, exp in enumerate(experience):
            title = exp.get("title", "")
            company = exp.get("company", "")
            period = exp.get("period", "")

            if idx > 0:
                _add_docx_job_separator(doc)

            header_parts = []
            if period:
                header_parts.append(period)
            if title:
                header_parts.append(title)
            if company:
                header_parts.append(company)
            if header_parts:
                _add_docx_job_header(doc, " | ".join(header_parts))

            for ach in exp.get("achievements", []):
                if ach.strip() and not _is_empty_content(ach):
                    _add_docx_bullet_paragraph(doc, f"• {ach}")

            honors = exp.get("honors", "")
            if honors and honors.strip() and not _is_empty_content(honors):
                _add_docx_bullet_paragraph(doc, f"★ {honors}")

    education = [e for e in cv_data.get("education", []) if _has_real_edu(e)]
    if education:
        _add_docx_section_header(doc, "השכלה")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            honors = edu.get("honors", "")
            parts = []
            if year:
                parts.append(year)
            if degree:
                parts.append(degree)
            if institution:
                parts.append(institution)
            if parts:
                _add_docx_body_paragraph(doc, " | ".join(parts))
            if honors and honors.strip() and not _is_empty_content(honors):
                _add_docx_bullet_paragraph(doc, f"★ {honors}")

    skills = cv_data.get("skills", {})
    technical = _filter_list(skills.get("technical", []))
    soft = _filter_list(skills.get("soft", []))
    if technical or soft:
        _add_docx_section_header(doc, "מיומנויות")
        if technical:
            _add_docx_body_paragraph(doc, ", ".join(technical))
        if soft:
            _add_docx_body_paragraph(doc, ", ".join(soft))

    lang_parts = []
    for lang in cv_data.get("languages", []):
        lang_name = lang.get("language", "").strip()
        level = lang.get("level", "").strip()
        if lang_name and not _is_empty_content(lang_name):
            part = lang_name
            if level and not _is_empty_content(level):
                part += f" – {level}"
            lang_parts.append(part)
    if lang_parts:
        _add_docx_section_header(doc, "שפות")
        _add_docx_body_paragraph(doc, " | ".join(lang_parts))

    military = _filter_list(cv_data.get("military", []))
    if military:
        _add_docx_section_header(doc, "שירות צבאי / לאומי")
        for item in military:
            _add_docx_bullet_paragraph(doc, f"• {item}")

    volunteering = _filter_list(cv_data.get("volunteering", []))
    if volunteering:
        _add_docx_section_header(doc, "התנדבות")
        for item in volunteering:
            _add_docx_bullet_paragraph(doc, f"• {item}")

    projects = _filter_list(cv_data.get("projects", []))
    if projects:
        _add_docx_section_header(doc, "פרויקטים עצמאיים")
        for item in projects:
            _add_docx_bullet_paragraph(doc, f"• {item}")

    additional = _filter_list(cv_data.get("additional", []))
    if additional:
        _add_docx_section_header(doc, "מידע נוסף")
        for item in additional:
            _add_docx_bullet_paragraph(doc, f"• {item}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _is_skills_section(title: str) -> bool:
    """Return True if the section title represents a skills/tools section."""
    keywords = ["מיומנויות", "כלים", "טכנולוגי", "skills", "tools", "technologies", "technical"]
    t = title.lower().strip()
    return any(k.lower() in t for k in keywords)


def _chunk_skills_lines(content: str, chunk_size: int = 4) -> list:
    """
    Split comma-separated skills content into bullet-ready chunks.
    If the content already uses bullet formatting, return lines as-is.
    Returns a list of strings, each representing one bullet line.
    """
    lines = [l.strip() for l in content.split("\n") if l.strip()]
    # If already bulleted — return as-is
    if any(l.startswith("•") or l.startswith("-") for l in lines):
        return lines
    # Collect all items from comma-separated lines
    all_items = []
    for line in lines:
        for item in line.split(","):
            item = item.strip()
            if item:
                all_items.append(item)
    if not all_items:
        return lines
    # Group into chunks
    chunks = []
    for i in range(0, len(all_items), chunk_size):
        chunk = all_items[i:i + chunk_size]
        chunks.append("• " + ", ".join(chunk))
    return chunks


def _is_job_header_line(line: str) -> bool:
    import re
    line = line.strip()
    if line.startswith("-") or line.startswith("•") or line.startswith("–"):
        return False
    military_keywords_he = ['שירות מלא', 'שירות סדיר', 'חיל ', 'צה"ל', 'צבא', 'שירות לאומי', 'שירות צבאי']
    for keyword in military_keywords_he:
        if keyword in line:
            return False
    lower = line.lower()
    military_keywords_en = [
        'military service', 'army', 'air force', 'navy', 'idf',
        'israeli defense', 'israeli air force', 'israeli navy',
        'national service', 'full service', 'combat', 'compulsory service'
    ]
    for keyword in military_keywords_en:
        if keyword in lower:
            return False
    if re.search(r'(19|20)\d{2}', line) and ('–' in line or '-' in line or '|' in line or 'הווה' in line or 'נוכחי' in line or 'היום' in line or 'present' in lower):
        return True
    return False


def _is_military_line(line: str) -> bool:
    lower = line.lower().strip()
    military_keywords = [
        'military service', 'military', 'idf', 'israeli defense', 'israeli air force',
        'israeli navy', 'combat', 'national service', 'army', 'navy', 'air force',
        'שירות צבאי', 'שירות סדיר', 'שירות מלא', 'שירות לאומי', 'צה"ל', 'צבא',
        'חיל הים', 'חיל האוויר',
    ]
    for keyword in military_keywords:
        if keyword in lower:
            return True
    return False


def _extract_contact_value(line: str) -> str:
    labels = ["טלפון", "אימייל", "מייל", "דוא\"ל", "לינקדאין", "עיר מגורים", "עיר", "כתובת", "אזור מגורים", "מגורים", "linkedin", "phone", "email", "city", "address", "residence"]
    val = line.strip()
    val_lower = val.lower()
    for label in labels:
        if val_lower.startswith(label.lower()):
            val = val[len(label):].strip()
            val = val.lstrip("-–—:").strip()
            break
    return val


def _set_docx_rtl(paragraph, font_name="Assistant"):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc_elem = pPr.find(qn('w:jc'))
    if jc_elem is not None:
        jc_val = jc_elem.get(qn('w:val'))
        if jc_val != 'center':
            pPr.remove(jc_elem)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:cs'), font_name)
        if run.font.bold:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)


def export_improved_cv_to_pdf(sections: list, cv_text: str = "", cv_title: str = "", max_pages: int = 1) -> bytes:
    font_name = register_hebrew_font()
    bold_font = f"{font_name}-Bold" if font_name == "Assistant" else "Helvetica-Bold"

    for level in range(len(_PDF_COMPRESSION_LEVELS)):
        cparams = _PDF_COMPRESSION_LEVELS[level]
        margin_mm = cparams["margin_mm"]

        buffer = io.BytesIO()
        top_mm = cparams.get("top_mm", 12)
        bot_mm = cparams.get("bot_mm", 12)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=margin_mm * mm,
            leftMargin=margin_mm * mm,
            topMargin=top_mm * mm,
            bottomMargin=bot_mm * mm
        )
        styles = _get_pdf_styles(font_name, bold_font, cparams)
        elements = []

        if cv_title and cv_title.strip():
            _title_style = ParagraphStyle(
                "CVTitle", fontName=bold_font, fontSize=14, leading=18,
                alignment=TA_CENTER, textColor=HexColor("#022559"), spaceAfter=4
            )
            elements.append(Paragraph(reshape_hebrew(cv_title.strip()), _title_style))

        for section in sections:
            title = section.get("title", "")
            content = section.get("final_text", section.get("improved", ""))
            is_personal = title and any(k in title for k in ["פרטים", "אישיים", "personal", "Personal"])
            if not is_personal and _is_empty_content(content):
                continue
            if title and not is_personal:
                elements.append(Paragraph(reshape_hebrew(title), styles["section_header"]))
                elements.append(_make_section_separator())
            if content:
                if is_personal:
                    lines = [l.strip() for l in content.split("\n") if l.strip()]
                    contact_only = [l for l in lines if not _is_military_line(l)]
                    if contact_only:
                        first_line = contact_only[0]
                        if "|" in first_line:
                            pipe_parts = [p.strip() for p in first_line.split("|")]
                            name = pipe_parts[0]
                            inline_contact = pipe_parts[1:]
                            extra_lines = contact_only[1:]
                        else:
                            name = first_line
                            inline_contact = []
                            extra_lines = contact_only[1:]
                        elements.append(Paragraph(reshape_hebrew(name), styles["name"]))
                        contact_values = []
                        for cl in inline_contact + extra_lines:
                            val = _extract_contact_value(cl)
                            if val:
                                contact_values.append(reshape_hebrew(val))
                        if contact_values:
                            elements.append(Paragraph(" | ".join(contact_values), styles["contact"]))
                        elements.append(HRFlowable(width="100%", thickness=2, color=HexColor("#2c3e50"), spaceAfter=4, spaceBefore=2))
                else:
                    render_lines = (
                        _chunk_skills_lines(content)
                        if _is_skills_section(title)
                        else [l.strip() for l in content.split("\n") if l.strip()]
                    )
                    for stripped in render_lines:
                        if not stripped:
                            continue
                        if _is_job_header_line(stripped):
                            elements.append(Paragraph(reshape_hebrew(stripped), styles["job_title"]))
                        elif stripped.startswith("-") or stripped.startswith("•"):
                            elements.append(Paragraph(reshape_hebrew(stripped), styles["bullet"]))
                        else:
                            elements.append(Paragraph(reshape_hebrew_paragraph(stripped), styles["body"]))

        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        if _count_pdf_pages(pdf_bytes) <= max_pages or level == len(_PDF_COMPRESSION_LEVELS) - 1:
            return pdf_bytes

    return buffer.getvalue()


def export_improved_cv_to_docx(sections: list, cv_text: str = "", cv_title: str = "") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Assistant"
    font.size = Pt(9)
    font.color.rgb = RGBColor(51, 51, 51)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(1)

    for s in doc.sections:
        s.right_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.top_margin = Cm(1.0)
        s.bottom_margin = Cm(1.0)

    if cv_title and cv_title.strip():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cv_title.strip())
        run.bold = True
        run.font.name = "Assistant"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(2, 37, 89)
        p.paragraph_format.space_after = Pt(4)

    for section in sections:
        title = section.get("title", "")
        content = section.get("final_text", section.get("improved", ""))
        is_personal = title and any(k in title for k in ["פרטים", "אישיים", "personal", "Personal"])

        if not is_personal and _is_empty_content(content):
            continue

        if title and not is_personal:
            _add_docx_section_header(doc, title)

        if content:
            if is_personal:
                lines = [l.strip() for l in content.split("\n") if l.strip()]
                contact_only = [l for l in lines if not _is_military_line(l)]
                if contact_only:
                    first_line = contact_only[0]
                    if "|" in first_line:
                        pipe_parts = [p.strip() for p in first_line.split("|")]
                        name_str = pipe_parts[0]
                        inline_contact = pipe_parts[1:]
                        extra_lines = contact_only[1:]
                    else:
                        name_str = first_line
                        inline_contact = []
                        extra_lines = contact_only[1:]
                    p = doc.add_paragraph()
                    run = p.add_run(name_str)
                    run.font.bold = True
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(44, 62, 80)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = Pt(24)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_docx_rtl(p)
                    contact_values = []
                    for cl in inline_contact + extra_lines:
                        val = _extract_contact_value(cl)
                        if val:
                            contact_values.append(val)
                    if contact_values:
                        p = doc.add_paragraph()
                        run = p.add_run(" | ".join(contact_values))
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(85, 85, 85)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_docx_rtl(p)
                    _add_docx_hr(doc)
            else:
                render_lines = (
                    _chunk_skills_lines(content)
                    if _is_skills_section(title)
                    else [l.strip() for l in content.split("\n") if l.strip()]
                )
                for stripped in render_lines:
                    if not stripped:
                        continue
                    if _is_job_header_line(stripped):
                        _add_docx_job_header(doc, stripped)
                    elif stripped.startswith("-") or stripped.startswith("•"):
                        _add_docx_bullet_paragraph(doc, stripped)
                    else:
                        _add_docx_body_paragraph(doc, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _get_pdf_styles_en(font_name, bold_font, cparams=None):
    from reportlab.lib.enums import TA_LEFT
    if cparams is None:
        cparams = _PDF_COMPRESSION_LEVELS[0]
    bs = cparams["body_size"]
    ss = cparams["sec_size"]
    bl = cparams["body_lead"]
    ul = cparams["bul_lead"]
    sb = cparams["sec_before"]
    ca = cparams["contact_after"]
    nl = cparams["name_lead"]
    return {
        "name": ParagraphStyle(
            "NameEN",
            fontName=bold_font,
            fontSize=20,
            leading=nl,
            alignment=TA_CENTER,
            textColor=HexColor("#2c3e50"),
            spaceAfter=2
        ),
        "contact": ParagraphStyle(
            "ContactEN",
            fontName=font_name,
            fontSize=bs,
            leading=ul,
            alignment=TA_CENTER,
            textColor=HexColor("#555555"),
            spaceAfter=ca
        ),
        "section_header": ParagraphStyle(
            "SectionHeaderEN",
            fontName=bold_font,
            fontSize=ss,
            leading=ss + 4,
            alignment=TA_LEFT,
            textColor=HexColor("#2c3e50"),
            spaceAfter=2,
            spaceBefore=sb
        ),
        "job_title": ParagraphStyle(
            "JobTitleEN",
            fontName=bold_font,
            fontSize=bs,
            leading=bl,
            alignment=TA_LEFT,
            textColor=HexColor("#333333"),
            spaceAfter=1,
            spaceBefore=max(int(sb / 2), 2)
        ),
        "body": ParagraphStyle(
            "BodyEN",
            fontName=font_name,
            fontSize=bs,
            leading=bl,
            alignment=TA_LEFT,
            textColor=HexColor("#333333"),
            spaceAfter=2
        ),
        "body_bold": ParagraphStyle(
            "BodyBoldEN",
            fontName=bold_font,
            fontSize=bs,
            leading=bl,
            alignment=TA_LEFT,
            textColor=HexColor("#2c3e50"),
            spaceAfter=1
        ),
        "bullet": ParagraphStyle(
            "BulletEN",
            fontName=font_name,
            fontSize=bs,
            leading=ul,
            alignment=TA_LEFT,
            textColor=HexColor("#333333"),
            spaceAfter=1,
            leftIndent=8
        ),
    }


def export_cv_to_pdf_en(cv_data: dict, max_pages: int = 1) -> bytes:
    font_name = register_hebrew_font()
    bold_font = f"{font_name}-Bold" if font_name == "Assistant" else "Helvetica-Bold"

    for level in range(len(_PDF_COMPRESSION_LEVELS)):
        cparams = _PDF_COMPRESSION_LEVELS[level]
        margin_mm = cparams["margin_mm"]

        buffer = io.BytesIO()
        top_mm = cparams.get("top_mm", 12)
        bot_mm = cparams.get("bot_mm", 12)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=margin_mm * mm,
            leftMargin=margin_mm * mm,
            topMargin=top_mm * mm,
            bottomMargin=bot_mm * mm
        )
        styles = _get_pdf_styles_en(font_name, bold_font, cparams)
        elements = []

        name = cv_data.get("full_name", "")
        if name:
            elements.append(Paragraph(name, styles["name"]))

        contact = cv_data.get("contact", {})
        contact_parts = []
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("city"):
            contact_parts.append(contact["city"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact_parts:
            elements.append(Paragraph(" | ".join(contact_parts), styles["contact"]))

        elements.append(HRFlowable(width="100%", thickness=2, color=HexColor("#2c3e50"), spaceAfter=4, spaceBefore=2))

        summary = cv_data.get("professional_summary", "")
        if summary and not _is_empty_content(summary):
            elements.append(Paragraph("Professional Summary", styles["section_header"]))
            elements.append(_make_section_separator())
            elements.append(Paragraph(summary, styles["body"]))

        experience = [e for e in cv_data.get("experience", []) if _has_real_exp(e)]
        if experience:
            elements.append(Paragraph("Professional Experience", styles["section_header"]))
            elements.append(_make_section_separator())
            for idx, exp in enumerate(experience):
                title = exp.get("title", "")
                company = exp.get("company", "")
                period = exp.get("period", "")
                if idx > 0:
                    elements.append(_make_job_separator())
                header_parts = []
                if period:
                    header_parts.append(period)
                if title:
                    header_parts.append(title)
                if company:
                    header_parts.append(company)
                if header_parts:
                    elements.append(Paragraph(" | ".join(header_parts), styles["job_title"]))
                for ach in exp.get("achievements", []):
                    if ach.strip() and not _is_empty_content(ach):
                        elements.append(Paragraph(f"• {ach}", styles["bullet"]))
                honors = exp.get("honors", "")
                if honors and honors.strip() and not _is_empty_content(honors):
                    elements.append(Paragraph(f"★ {honors}", styles["bullet"]))

        education = [e for e in cv_data.get("education", []) if _has_real_edu(e)]
        if education:
            elements.append(Paragraph("Education", styles["section_header"]))
            elements.append(_make_section_separator())
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")
                honors = edu.get("honors", "")
                parts = []
                if year:
                    parts.append(year)
                if degree:
                    parts.append(degree)
                if institution:
                    parts.append(institution)
                if parts:
                    elements.append(Paragraph(" | ".join(parts), styles["body"]))
                if honors and honors.strip() and not _is_empty_content(honors):
                    elements.append(Paragraph(f"★ {honors}", styles["bullet"]))

        skills = cv_data.get("skills", {})
        technical = _filter_list(skills.get("technical", []))
        soft = _filter_list(skills.get("soft", []))
        if technical or soft:
            elements.append(Paragraph("Skills", styles["section_header"]))
            elements.append(_make_section_separator())
            if technical:
                elements.append(Paragraph(", ".join(technical), styles["body"]))
            if soft:
                elements.append(Paragraph(", ".join(soft), styles["body"]))

        lang_parts = []
        for lang in cv_data.get("languages", []):
            lang_name = lang.get("language", "").strip()
            lang_level = lang.get("level", "").strip()
            if lang_name and not _is_empty_content(lang_name):
                part = lang_name
                if lang_level and not _is_empty_content(lang_level):
                    part += f" – {lang_level}"
                lang_parts.append(part)
        if lang_parts:
            elements.append(Paragraph("Languages", styles["section_header"]))
            elements.append(_make_section_separator())
            elements.append(Paragraph(" | ".join(lang_parts), styles["body"]))

        military = _filter_list(cv_data.get("military", []))
        if military:
            elements.append(Paragraph("Military / National Service", styles["section_header"]))
            elements.append(_make_section_separator())
            for item in military:
                elements.append(Paragraph(f"• {item}", styles["bullet"]))

        volunteering = _filter_list(cv_data.get("volunteering", []))
        if volunteering:
            elements.append(Paragraph("Volunteering", styles["section_header"]))
            elements.append(_make_section_separator())
            for item in volunteering:
                elements.append(Paragraph(f"• {item}", styles["bullet"]))

        projects = _filter_list(cv_data.get("projects", []))
        if projects:
            elements.append(Paragraph("Personal Projects", styles["section_header"]))
            elements.append(_make_section_separator())
            for item in projects:
                elements.append(Paragraph(f"• {item}", styles["bullet"]))

        additional = _filter_list(cv_data.get("additional", []))
        if additional:
            elements.append(Paragraph("Additional Information", styles["section_header"]))
            elements.append(_make_section_separator())
            for item in additional:
                elements.append(Paragraph(f"• {item}", styles["bullet"]))

        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        if _count_pdf_pages(pdf_bytes) <= max_pages or level == len(_PDF_COMPRESSION_LEVELS) - 1:
            return pdf_bytes

    return buffer.getvalue()


def export_cv_to_docx_en(cv_data: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Assistant"
    font.size = Pt(9)
    font.color.rgb = RGBColor(51, 51, 51)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(1)

    for section in doc.sections:
        section.right_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    name = cv_data.get("full_name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(24)

    contact = cv_data.get("contact", {})
    contact_parts = []
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("city"):
        contact_parts.append(contact["city"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(12)

    _add_docx_hr(doc)

    summary = cv_data.get("professional_summary", "")
    if summary and not _is_empty_content(summary):
        _add_docx_section_header_en(doc, "Professional Summary")
        _add_docx_body_paragraph(doc, summary, is_rtl=False)

    experience = [e for e in cv_data.get("experience", []) if _has_real_exp(e)]
    if experience:
        _add_docx_section_header_en(doc, "Professional Experience")
        for idx, exp in enumerate(experience):
            title = exp.get("title", "")
            company = exp.get("company", "")
            period = exp.get("period", "")
            if idx > 0:
                _add_docx_job_separator(doc)
            header_parts = []
            if period:
                header_parts.append(period)
            if title:
                header_parts.append(title)
            if company:
                header_parts.append(company)
            if header_parts:
                _add_docx_job_header(doc, " | ".join(header_parts), is_rtl=False)
            for ach in exp.get("achievements", []):
                if ach.strip() and not _is_empty_content(ach):
                    _add_docx_bullet_paragraph(doc, f"• {ach}", is_rtl=False)
            honors = exp.get("honors", "")
            if honors and honors.strip() and not _is_empty_content(honors):
                _add_docx_bullet_paragraph(doc, f"★ {honors}", is_rtl=False)

    education = [e for e in cv_data.get("education", []) if _has_real_edu(e)]
    if education:
        _add_docx_section_header_en(doc, "Education")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            honors = edu.get("honors", "")
            parts = []
            if year:
                parts.append(year)
            if degree:
                parts.append(degree)
            if institution:
                parts.append(institution)
            if parts:
                _add_docx_body_paragraph(doc, " | ".join(parts), is_rtl=False)
            if honors and honors.strip() and not _is_empty_content(honors):
                _add_docx_bullet_paragraph(doc, f"★ {honors}", is_rtl=False)

    skills = cv_data.get("skills", {})
    technical = _filter_list(skills.get("technical", []))
    soft = _filter_list(skills.get("soft", []))
    if technical or soft:
        _add_docx_section_header_en(doc, "Skills")
        if technical:
            _add_docx_body_paragraph(doc, ", ".join(technical), is_rtl=False)
        if soft:
            _add_docx_body_paragraph(doc, ", ".join(soft), is_rtl=False)

    lang_parts = []
    for lang in cv_data.get("languages", []):
        lang_name = lang.get("language", "").strip()
        level = lang.get("level", "").strip()
        if lang_name and not _is_empty_content(lang_name):
            part = lang_name
            if level and not _is_empty_content(level):
                part += f" – {level}"
            lang_parts.append(part)
    if lang_parts:
        _add_docx_section_header_en(doc, "Languages")
        _add_docx_body_paragraph(doc, " | ".join(lang_parts), is_rtl=False)

    military = _filter_list(cv_data.get("military", []))
    if military:
        _add_docx_section_header_en(doc, "Military / National Service")
        for item in military:
            _add_docx_bullet_paragraph(doc, f"• {item}", is_rtl=False)

    volunteering = _filter_list(cv_data.get("volunteering", []))
    if volunteering:
        _add_docx_section_header_en(doc, "Volunteering")
        for item in volunteering:
            _add_docx_bullet_paragraph(doc, f"• {item}", is_rtl=False)

    projects = _filter_list(cv_data.get("projects", []))
    if projects:
        _add_docx_section_header_en(doc, "Personal Projects")
        for item in projects:
            _add_docx_bullet_paragraph(doc, f"• {item}", is_rtl=False)

    additional = _filter_list(cv_data.get("additional", []))
    if additional:
        _add_docx_section_header_en(doc, "Additional Information")
        for item in additional:
            _add_docx_bullet_paragraph(doc, f"• {item}", is_rtl=False)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_docx_section_header_en(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)

    _add_docx_separator_line(doc, '7fb3d8', '8', 0, 1)


def _is_section_header_en(line: str) -> bool:
    import re
    stripped = line.strip()
    section_keywords = [
        "personal details", "professional summary", "professional experience",
        "work experience", "education", "skills", "languages", "additional",
        "contact", "certifications", "training", "volunteering", "military",
        "profile", "objective", "qualifications", "additional information",
        "summary", "experience", "personal information", "contact details",
        "contact information"
    ]
    if re.match(r'^===\s*(.+?)\s*===$', stripped):
        return True
    lower = stripped.lower().rstrip(":")
    if lower in section_keywords:
        return True
    if len(stripped) < 40 and not stripped.startswith("•") and not stripped.startswith("-") and stripped.endswith(":"):
        return True
    return False


def _clean_section_title(line: str) -> str:
    import re
    stripped = line.strip()
    m = re.match(r'^===\s*(.+?)\s*===$', stripped)
    if m:
        return m.group(1).rstrip(":")
    return stripped.rstrip(":")


def export_improved_cv_to_pdf_en(translated_text: str, cv_title: str = "", max_pages: int = 1) -> bytes:
    font_name = register_hebrew_font()
    bold_font = f"{font_name}-Bold" if font_name == "Assistant" else "Helvetica-Bold"

    for level in range(len(_PDF_COMPRESSION_LEVELS)):
        cparams = _PDF_COMPRESSION_LEVELS[level]
        margin_mm = cparams["margin_mm"]

        buffer = io.BytesIO()
        top_mm = cparams.get("top_mm", 12)
        bot_mm = cparams.get("bot_mm", 12)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=margin_mm * mm,
            leftMargin=margin_mm * mm,
            topMargin=top_mm * mm,
            bottomMargin=bot_mm * mm
        )
        styles = _get_pdf_styles_en(font_name, bold_font, cparams)
        elements = []

        if cv_title and cv_title.strip():
            from reportlab.lib.enums import TA_CENTER as _TA_CENTER
            _title_style = ParagraphStyle(
                "CVTitleEn", fontName=bold_font, fontSize=14, leading=18,
                alignment=_TA_CENTER, textColor=HexColor("#2c3e50"), spaceAfter=4
            )
            elements.append(Paragraph(cv_title.strip(), _title_style))

        last_header = None
        in_personal = False
        personal_lines = []
        deferred_military_lines = []
        pending_header = None  # deferred — only emitted when real content follows
        for line in translated_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if _is_section_header_en(stripped):
                if in_personal and personal_lines:
                    contact_only = [l for l in personal_lines if not _is_military_line(l)]
                    military_only = [l for l in personal_lines if _is_military_line(l)]
                    deferred_military_lines.extend(military_only)
                    if contact_only:
                        elements.append(Paragraph(contact_only[0], styles["name"]))
                        contact_values = []
                        for cl in contact_only[1:]:
                            val = _extract_contact_value(cl)
                            if val:
                                contact_values.append(val)
                        if contact_values:
                            elements.append(Paragraph(" | ".join(contact_values), styles["contact"]))
                    elements.append(HRFlowable(width="100%", thickness=2, color=HexColor("#2c3e50"), spaceAfter=4, spaceBefore=2))
                    personal_lines = []
                header_text = _clean_section_title(stripped)
                if header_text == last_header:
                    continue
                lower_header = header_text.lower()
                if any(k in lower_header for k in ["personal", "contact"]):
                    in_personal = True
                    last_header = header_text
                    pending_header = None
                    continue
                in_personal = False
                last_header = header_text
                pending_header = header_text  # defer; emit only when real content follows
            elif in_personal:
                personal_lines.append(stripped)
            elif _is_job_header_line(stripped):
                if not _is_empty_content(stripped):
                    if pending_header is not None:
                        elements.append(Paragraph(pending_header, styles["section_header"]))
                        elements.append(_make_section_separator())
                        pending_header = None
                    elements.append(Paragraph(stripped, styles["job_title"]))
            elif stripped.startswith("-") or stripped.startswith("•"):
                if not _is_empty_content(stripped):
                    if pending_header is not None:
                        elements.append(Paragraph(pending_header, styles["section_header"]))
                        elements.append(_make_section_separator())
                        pending_header = None
                    elements.append(Paragraph(stripped, styles["bullet"]))
            else:
                if not _is_empty_content(stripped):
                    if pending_header is not None:
                        elements.append(Paragraph(pending_header, styles["section_header"]))
                        elements.append(_make_section_separator())
                        pending_header = None
                    elements.append(Paragraph(stripped, styles["body"]))

        if in_personal and personal_lines:
            contact_only = [l for l in personal_lines if not _is_military_line(l)]
            military_only = [l for l in personal_lines if _is_military_line(l)]
            deferred_military_lines.extend(military_only)
            if contact_only:
                elements.append(Paragraph(contact_only[0], styles["name"]))
                contact_values = []
                for cl in contact_only[1:]:
                    val = _extract_contact_value(cl)
                    if val:
                        contact_values.append(val)
                if contact_values:
                    elements.append(Paragraph(" | ".join(contact_values), styles["contact"]))
            elements.append(HRFlowable(width="100%", thickness=2, color=HexColor("#2c3e50"), spaceAfter=4, spaceBefore=2))

        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        if _count_pdf_pages(pdf_bytes) <= max_pages or level == len(_PDF_COMPRESSION_LEVELS) - 1:
            return pdf_bytes

    return buffer.getvalue()


def _add_docx_personal_block_en(doc, personal_lines):
    if not personal_lines:
        return []
    contact_only = [l for l in personal_lines if not _is_military_line(l)]
    military_only = [l for l in personal_lines if _is_military_line(l)]
    if not contact_only:
        _add_docx_hr(doc)
        return military_only
    p = doc.add_paragraph()
    run = p.add_run(contact_only[0])
    run.font.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(44, 62, 80)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_values = []
    for cl in contact_only[1:]:
        val = _extract_contact_value(cl)
        if val:
            contact_values.append(val)
    if contact_values:
        p = doc.add_paragraph()
        run = p.add_run(" | ".join(contact_values))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_docx_hr(doc)
    return military_only


def export_improved_cv_to_docx_en(translated_text: str, cv_title: str = "") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Assistant"
    font.size = Pt(9)
    font.color.rgb = RGBColor(51, 51, 51)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(1)

    for s in doc.sections:
        s.right_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.top_margin = Cm(1.0)
        s.bottom_margin = Cm(1.0)

    if cv_title and cv_title.strip():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cv_title.strip())
        run.bold = True
        run.font.name = "Assistant"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_after = Pt(4)

    last_header = None
    in_personal = False
    personal_lines = []
    deferred_military_lines = []
    pending_header = None  # deferred — only emitted when real content follows
    for line in translated_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if _is_section_header_en(stripped):
            if in_personal and personal_lines:
                mil = _add_docx_personal_block_en(doc, personal_lines)
                deferred_military_lines.extend(mil)
                personal_lines = []
            header_text = _clean_section_title(stripped)
            if header_text == last_header:
                continue
            lower_header = header_text.lower()
            if any(k in lower_header for k in ["personal", "contact"]):
                in_personal = True
                last_header = header_text
                pending_header = None
                continue
            in_personal = False
            last_header = header_text
            pending_header = header_text  # defer; emit only when real content follows
        elif in_personal:
            personal_lines.append(stripped)
        elif _is_job_header_line(stripped):
            if not _is_empty_content(stripped):
                if pending_header is not None:
                    _add_docx_section_header_en(doc, pending_header)
                    pending_header = None
                _add_docx_job_header(doc, stripped, is_rtl=False)
        elif stripped.startswith("-") or stripped.startswith("•"):
            if not _is_empty_content(stripped):
                if pending_header is not None:
                    _add_docx_section_header_en(doc, pending_header)
                    pending_header = None
                _add_docx_bullet_paragraph(doc, stripped, is_rtl=False)
        else:
            if not _is_empty_content(stripped):
                if pending_header is not None:
                    _add_docx_section_header_en(doc, pending_header)
                    pending_header = None
                _add_docx_body_paragraph(doc, stripped, is_rtl=False)

    if in_personal and personal_lines:
        mil = _add_docx_personal_block_en(doc, personal_lines)
        deferred_military_lines.extend(mil)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
